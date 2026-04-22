"""
Script para generar el documento de metodología de desarrollo del sistema
"Transcription & Analysis" - Frontend + Backend completo.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ─── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Apply background color to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table_row(table, cells_data, bold=False, header=False, bg_color=None):
    """Add a row to a table with formatting."""
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(9)
        if bold or header:
            run.bold = True
        if header:
            run.font.color.rgb = RGBColor(255, 255, 255)
            set_cell_shading(cell, "2B579A")
        elif bg_color:
            set_cell_shading(cell, bg_color)
    return row

def create_table(doc, headers, rows, col_widths=None):
    """Create a formatted table."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, "2B579A")

    # Data rows
    for idx, row_data in enumerate(rows):
        bg = "F2F2F2" if idx % 2 == 1 else None
        add_table_row(table, row_data, bg_color=bg)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    return table

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(43, 87, 154)
    return h

# ─── MAIN DOCUMENT ───────────────────────────────────────────────────────────

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)

# ══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("METODOLOGÍA DE DESARROLLO DE SOFTWARE")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(43, 87, 154)

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Sistema de Transcripción y Análisis de Audios")
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(89, 89, 89)

doc.add_paragraph()

subtitle2 = doc.add_paragraph()
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle2.add_run("Transcription & Analysis Platform")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(128, 128, 128)
run.italic = True

doc.add_paragraph()
doc.add_paragraph()

# Info table portada
info_table = doc.add_table(rows=6, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_data = [
    ("Documento", "Metodología de Desarrollo del Sistema"),
    ("Versión", "1.0"),
    ("Fecha", "Abril 2026"),
    ("Estado", "En Desarrollo"),
    ("Backend", "Python - FastAPI"),
    ("Frontend", "Angular 21 - TypeScript"),
]
for i, (label, value) in enumerate(info_data):
    cell_l = info_table.rows[i].cells[0]
    cell_l.text = ""
    r = cell_l.paragraphs[0].add_run(label)
    r.bold = True
    r.font.size = Pt(11)
    set_cell_shading(cell_l, "E8EEF7")

    cell_r = info_table.rows[i].cells[1]
    cell_r.text = ""
    r = cell_r.paragraphs[0].add_run(value)
    r.font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# ÍNDICE
# ══════════════════════════════════════════════════════════════════════════════

add_heading("ÍNDICE DE CONTENIDOS", level=1)

index_items = [
    "1. Introducción",
    "   1.1. Propósito del Documento",
    "   1.2. Alcance del Sistema",
    "   1.3. Objetivos del Proyecto",
    "2. Descripción General del Sistema",
    "   2.1. Perspectiva del Producto",
    "   2.2. Funcionalidades Principales",
    "   2.3. Usuarios del Sistema",
    "3. Stack Tecnológico",
    "   3.1. Backend - Python/FastAPI",
    "   3.2. Frontend - Angular 21",
    "   3.3. Base de Datos - PostgreSQL",
    "   3.4. Servicios Externos",
    "4. Arquitectura del Sistema",
    "   4.1. Arquitectura General (Cliente-Servidor)",
    "   4.2. Arquitectura del Backend",
    "   4.3. Arquitectura del Frontend",
    "   4.4. Modelo de Datos",
    "5. Tabla de Requerimientos Funcionales",
    "6. Tabla de Requerimientos No Funcionales",
    "7. Especificación de Módulos",
    "   7.1. Módulo de Autenticación",
    "   7.2. Módulo de Campañas",
    "   7.3. Módulo de Audios",
    "   7.4. Módulo de Análisis",
    "   7.5. Módulo de Configuración",
    "   7.6. Módulo Dashboard",
    "8. Especificación de Endpoints API",
    "   8.1. Auth API",
    "   8.2. Campaigns API",
    "   8.3. Audios API",
    "   8.4. Analysis API",
    "   8.5. Config API",
    "9. Estructura del Proyecto",
    "   9.1. Estructura Backend",
    "   9.2. Estructura Frontend",
    "10. Flujos de Proceso Principales",
    "   10.1. Pipeline de Transcripción",
    "   10.2. Pipeline de Análisis con IA",
    "   10.3. Flujo de Autenticación",
    "11. Diseño de la Base de Datos",
    "12. Gestión de Configuración y Entorno",
    "13. Seguridad del Sistema",
    "14. Exportación de Datos",
    "15. Gestión del Proyecto",
    "   15.1. Metodología de Desarrollo",
    "   15.2. Control de Versiones",
    "   15.3. Convenciones de Código",
    "16. Plan de Pruebas",
    "17. Glosario de Términos",
]

for item in index_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    for run in p.runs:
        run.font.size = Pt(10)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCCIÓN
# ══════════════════════════════════════════════════════════════════════════════

add_heading("1. INTRODUCCIÓN", level=1)

add_heading("1.1. Propósito del Documento", level=2)
doc.add_paragraph(
    "El presente documento establece la metodología de desarrollo, los requerimientos funcionales "
    "y no funcionales, la arquitectura del sistema, y la especificación técnica completa del "
    "proyecto \"Sistema de Transcripción y Análisis de Audios\" (Transcription & Analysis Platform). "
    "Este documento sirve como guía integral para el equipo de desarrollo, documentando las decisiones "
    "técnicas, los patrones de diseño adoptados y las especificaciones de cada módulo tanto del "
    "backend como del frontend."
)

add_heading("1.2. Alcance del Sistema", level=2)
doc.add_paragraph(
    "El sistema es una plataforma web full-stack diseñada para la transcripción automatizada de "
    "audios de llamadas telefónicas y su posterior análisis de calidad mediante inteligencia "
    "artificial. Permite a los usuarios gestionar campañas, subir archivos de audio, transcribirlos "
    "usando servicios de reconocimiento de voz (Deepgram, WhisperX), y evaluar su contenido "
    "contra criterios de calidad personalizables utilizando modelos de lenguaje (OpenAI GPT)."
)

doc.add_paragraph("El sistema abarca los siguientes componentes:")
components = [
    "Backend API RESTful desarrollado con Python y FastAPI",
    "Frontend SPA (Single Page Application) desarrollado con Angular 21 y TypeScript",
    "Base de datos PostgreSQL 16 para persistencia",
    "Integración con servicios externos: Deepgram, WhisperX, OpenAI, Pyannote",
    "Sistema de exportación de resultados en múltiples formatos (XLSX, CSV, JSON)",
]
for c in components:
    doc.add_paragraph(c, style="List Bullet")

add_heading("1.3. Objetivos del Proyecto", level=2)
objectives = [
    ("OBJ-01", "Automatizar la transcripción de audios de llamadas telefónicas con identificación de hablantes (diarización)"),
    ("OBJ-02", "Proporcionar análisis de calidad automatizado mediante IA con criterios configurables"),
    ("OBJ-03", "Ofrecer una interfaz web intuitiva para la gestión integral de campañas, audios y análisis"),
    ("OBJ-04", "Permitir la exportación de resultados en formatos estándar para reportería"),
    ("OBJ-05", "Garantizar la seguridad de los datos mediante autenticación JWT y control de acceso por usuario"),
    ("OBJ-06", "Soportar procesamiento masivo (batch) de audios con monitoreo de progreso en tiempo real"),
]
create_table(doc, ["ID", "Objetivo"], objectives, col_widths=[2.5, 14])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2. DESCRIPCIÓN GENERAL DEL SISTEMA
# ══════════════════════════════════════════════════════════════════════════════

add_heading("2. DESCRIPCIÓN GENERAL DEL SISTEMA", level=1)

add_heading("2.1. Perspectiva del Producto", level=2)
doc.add_paragraph(
    "El sistema opera como una aplicación web cliente-servidor. El backend expone una API RESTful "
    "que gestiona la lógica de negocio, la persistencia de datos y la orquestación de servicios "
    "externos de transcripción y análisis. El frontend consume esta API proporcionando una "
    "interfaz gráfica moderna y responsiva para los usuarios finales."
)

add_heading("2.2. Funcionalidades Principales", level=2)
features = [
    ("Autenticación y Autorización", "Registro, login, gestión de perfil con JWT. Control de acceso por usuario."),
    ("Gestión de Campañas", "CRUD completo de campañas de llamadas. Estadísticas por campaña."),
    ("Gestión de Audios", "Carga unitaria y masiva de archivos de audio. Soporte drag-and-drop. Filtrado por campaña."),
    ("Transcripción Automatizada", "Transcripción con Deepgram (cloud) o WhisperX (local). Diarización de hablantes. Procesamiento batch."),
    ("Análisis con IA", "Evaluación de calidad con OpenAI GPT. Prompts personalizables. Formato de salida configurable."),
    ("Pipeline Integrado", "Ejecución secuencial automática: Transcripción → Análisis. Monitoreo en tiempo real."),
    ("Dashboard", "Estadísticas generales: total de audios, progreso de transcripción, costos, campañas activas."),
    ("Exportación de Datos", "Exportación en XLSX, CSV y JSON con configuración de columnas y metadatos."),
    ("Gestión de Plantillas", "CRUD de prompts de análisis y formatos de salida reutilizables."),
    ("Tema Visual", "Soporte light/dark mode con persistencia de preferencia del usuario."),
]
create_table(doc, ["Funcionalidad", "Descripción"], features, col_widths=[4, 12.5])

add_heading("2.3. Usuarios del Sistema", level=2)
users = [
    ("Analista QA", "Usuario principal. Gestiona campañas, sube audios, configura análisis y revisa resultados."),
    ("Supervisor", "Revisa dashboards y exporta reportes consolidados de calidad."),
    ("Administrador", "Gestiona usuarios y configuraciones globales del sistema."),
]
create_table(doc, ["Rol", "Descripción"], users, col_widths=[4, 12.5])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. STACK TECNOLÓGICO
# ══════════════════════════════════════════════════════════════════════════════

add_heading("3. STACK TECNOLÓGICO", level=1)

add_heading("3.1. Backend - Python/FastAPI", level=2)
backend_stack = [
    ("Lenguaje", "Python 3.14"),
    ("Framework Web", "FastAPI 0.125.0"),
    ("Servidor ASGI", "Uvicorn 0.38.0"),
    ("ORM", "SQLAlchemy 2.0.45"),
    ("Migraciones", "Alembic 1.17.2"),
    ("Validación", "Pydantic 2.12.5"),
    ("Autenticación", "python-jose 3.5.0 (JWT), passlib 1.7.4 (bcrypt)"),
    ("Procesamiento Audio", "Pydub 0.25.1, FFmpeg"),
    ("Machine Learning", "PyTorch, Transformers (HuggingFace)"),
    ("Exportación", "Pandas, Openpyxl 3.1.5"),
    ("Contenedores", "Docker Compose (PostgreSQL)"),
]
create_table(doc, ["Componente", "Tecnología / Versión"], backend_stack, col_widths=[4, 12.5])

add_heading("3.2. Frontend - Angular 21", level=2)
frontend_stack = [
    ("Framework", "Angular 21.0.0"),
    ("Lenguaje", "TypeScript 5.9.2 (strict mode)"),
    ("Reactividad", "RxJS 7.8.0, Angular Signals"),
    ("CSS Framework", "TailwindCSS 4.2.1"),
    ("UI Components", "DaisyUI 5.5.19"),
    ("Build Tool", "Angular CLI 21.0.3"),
    ("Testing", "Vitest 4.0.8, JSDOM 27.1.0"),
    ("Formatter", "Prettier (100 chars, single quotes)"),
    ("HTTP Client", "Angular HttpClient (fetch-based)"),
]
create_table(doc, ["Componente", "Tecnología / Versión"], frontend_stack, col_widths=[4, 12.5])

add_heading("3.3. Base de Datos - PostgreSQL", level=2)
db_stack = [
    ("Motor", "PostgreSQL 16"),
    ("Deployment", "Docker Compose"),
    ("ORM", "SQLAlchemy 2.0 (Declarative Mapped)"),
    ("Migraciones", "Alembic"),
    ("Tipos de datos clave", "UUID (PK), Text, Numeric(10,4), TIMESTAMP"),
]
create_table(doc, ["Componente", "Detalle"], db_stack, col_widths=[4, 12.5])

add_heading("3.4. Servicios Externos", level=2)
external = [
    ("Deepgram", "Transcripción cloud (modelo nova-3). Diarización de hablantes. Múltiples API keys con failover."),
    ("WhisperX", "Transcripción local (modelos tiny→large). Alineación temporal. GPU/CPU."),
    ("Pyannote", "Diarización de hablantes (modelo speaker-diarization-3.1). Requiere token HuggingFace."),
    ("OpenAI", "Análisis QA con GPT-4.1-mini. Salida JSON estructurada. Control de tokens y costos."),
    ("FFmpeg", "Conversión de formatos de audio (MP3, WAV). Extracción de duración."),
]
create_table(doc, ["Servicio", "Descripción"], external, col_widths=[3.5, 13])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. ARQUITECTURA DEL SISTEMA
# ══════════════════════════════════════════════════════════════════════════════

add_heading("4. ARQUITECTURA DEL SISTEMA", level=1)

add_heading("4.1. Arquitectura General (Cliente-Servidor)", level=2)
doc.add_paragraph(
    "El sistema sigue una arquitectura Cliente-Servidor de tres capas: Presentación (Frontend Angular), "
    "Lógica de Negocio (Backend FastAPI), y Persistencia (PostgreSQL). La comunicación entre "
    "frontend y backend se realiza mediante HTTP REST con autenticación Bearer JWT."
)

arch_layers = [
    ("Capa de Presentación", "Angular 21 SPA", "Interfaces de usuario, validación de formularios, estado local con Signals, caché en memoria, temas visuales."),
    ("Capa de API / Lógica", "FastAPI (Python)", "Endpoints REST, validación con Pydantic, orquestación de servicios, procesamiento background, autenticación JWT."),
    ("Capa de Servicios", "Servicios Python", "Transcripción (Deepgram/WhisperX), Análisis (OpenAI), Almacenamiento de archivos, Exportación de datos."),
    ("Capa de Datos", "PostgreSQL 16", "Persistencia relacional, 6 tablas principales, relaciones con CASCADE, UUIDs como PKs."),
]
create_table(doc, ["Capa", "Tecnología", "Responsabilidad"], arch_layers, col_widths=[3.5, 3.5, 9.5])

add_heading("4.2. Arquitectura del Backend", level=2)
doc.add_paragraph(
    "El backend sigue un patrón de arquitectura en capas (Layered Architecture) con separación clara "
    "entre Routers (controladores), Services (lógica de negocio/repositorios), y Models (ORM). "
    "Utiliza inyección de dependencias nativa de FastAPI para gestionar sesiones de BD y autenticación."
)

backend_layers = [
    ("Routers", "app/routers/", "Definición de endpoints HTTP. Validación de entrada con Pydantic. Manejo de respuestas HTTP."),
    ("Services", "app/services/", "Lógica de negocio. Repositorios de datos (patrón Repository). Integración con servicios externos."),
    ("Models", "app/components/models.py", "Modelos ORM SQLAlchemy. Definición de tablas y relaciones."),
    ("Schemas", "app/components/schemas.py", "DTOs Pydantic para request/response. Validación de datos."),
    ("Dependencies", "app/dependencies/", "Middleware de autenticación. Inyección de sesión de BD."),
    ("Utils", "app/utils/", "Funciones utilitarias: seguridad, archivos, costos, audio, fechas."),
]
create_table(doc, ["Capa", "Ubicación", "Responsabilidad"], backend_layers, col_widths=[3, 4.5, 9])

add_heading("4.3. Arquitectura del Frontend", level=2)
doc.add_paragraph(
    "El frontend sigue una arquitectura modular basada en features con componentes standalone "
    "(sin NgModules). Utiliza lazy loading para optimizar el rendimiento y Angular Signals "
    "combinado con RxJS para la gestión de estado reactivo."
)

frontend_layers = [
    ("Core", "src/app/core/", "Servicios singleton: configuración API, almacenamiento de tokens, caché, tema visual."),
    ("Features", "src/app/feactures/", "Módulos de negocio: auth, campaigns, audios, analysis, shell. Cada uno con pages, components, services."),
    ("Shared", "src/app/shared/", "Componentes reutilizables: navbar, sidebar, footer, loading, theme-buttons, form utilities."),
    ("API", "src/app/api/", "Interfaces TypeScript para tipado de request/response de la API."),
    ("Guards", "src/app/guards/", "Guards de ruta: AuthGuard (protege rutas autenticadas), GuestGuard (solo invitados)."),
]
create_table(doc, ["Capa", "Ubicación", "Responsabilidad"], frontend_layers, col_widths=[2.5, 4, 10])

add_heading("4.4. Modelo de Datos (Resumen)", level=2)
doc.add_paragraph(
    "La base de datos consta de 6 tablas principales con relaciones uno-a-muchos y eliminación "
    "en cascada. Todas las tablas usan UUID como clave primaria y timestamps automáticos."
)

data_model = [
    ("tbl_users", "Usuarios del sistema. Username único, password hasheado con bcrypt."),
    ("tbl_campaigns", "Campañas de llamadas. Pertenecen a un usuario (FK user_id)."),
    ("tbl_audios", "Archivos de audio. Pertenecen a una campaña (FK campaign_id). Almacenan transcripción, costo y duración."),
    ("tbl_audios_analysis", "Resultados de análisis QA por audio (FK audio_id). Criterio, evaluación, justificación, tokens, costo."),
    ("tbl_prompt_templates", "Plantillas de prompts para análisis IA. Pertenecen a un usuario (FK user_id)."),
    ("tbl_output_formats", "Formatos de salida configurables con campos JSON. Pertenecen a un usuario (FK user_id)."),
]
create_table(doc, ["Tabla", "Descripción"], data_model, col_widths=[4, 12.5])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. TABLA DE REQUERIMIENTOS FUNCIONALES
# ══════════════════════════════════════════════════════════════════════════════

add_heading("5. TABLA DE REQUERIMIENTOS FUNCIONALES", level=1)

doc.add_paragraph(
    "A continuación se detallan los requerimientos funcionales del sistema, organizados por módulo."
)

reqs_func = [
    # Auth
    ("RF-001", "Autenticación", "Registro de usuarios", "El sistema debe permitir el registro de nuevos usuarios con username y password.", "Alta", "Implementado"),
    ("RF-002", "Autenticación", "Inicio de sesión", "El sistema debe autenticar usuarios mediante username/password y retornar un token JWT.", "Alta", "Implementado"),
    ("RF-003", "Autenticación", "Verificación de sesión", "El sistema debe verificar la validez del token JWT en cada petición protegida.", "Alta", "Implementado"),
    ("RF-004", "Autenticación", "Cierre de sesión", "El sistema debe permitir al usuario cerrar sesión eliminando el token del almacenamiento.", "Alta", "Implementado"),
    ("RF-005", "Autenticación", "Perfil de usuario", "El sistema debe permitir consultar y actualizar los datos del perfil del usuario autenticado.", "Media", "Implementado"),
    # Campaigns
    ("RF-006", "Campañas", "Crear campaña", "El sistema debe permitir crear campañas con nombre y descripción opcional.", "Alta", "Implementado"),
    ("RF-007", "Campañas", "Listar campañas", "El sistema debe listar campañas del usuario con paginación y búsqueda.", "Alta", "Implementado"),
    ("RF-008", "Campañas", "Editar campaña", "El sistema debe permitir actualizar el nombre y descripción de una campaña.", "Alta", "Implementado"),
    ("RF-009", "Campañas", "Eliminar campaña", "El sistema debe eliminar la campaña y todos sus audios/análisis asociados (CASCADE).", "Alta", "Implementado"),
    ("RF-010", "Campañas", "Estadísticas de campaña", "El sistema debe mostrar estadísticas: total audios, transcritos, costos, duración.", "Media", "Implementado"),
    ("RF-011", "Campañas", "Top 100 campañas", "El sistema debe proporcionar un listado rápido de las últimas 100 campañas.", "Baja", "Implementado"),
    # Audios
    ("RF-012", "Audios", "Carga unitaria de audio", "El sistema debe permitir subir un archivo de audio asociado a una campaña.", "Alta", "Implementado"),
    ("RF-013", "Audios", "Carga masiva de audios", "El sistema debe permitir subir múltiples archivos de audio simultáneamente.", "Alta", "Implementado"),
    ("RF-014", "Audios", "Drag and drop", "El frontend debe soportar carga de archivos mediante arrastrar y soltar.", "Media", "Implementado"),
    ("RF-015", "Audios", "Listar audios", "El sistema debe listar audios con paginación, búsqueda y filtro por campaña.", "Alta", "Implementado"),
    ("RF-016", "Audios", "Editar audio", "El sistema debe permitir actualizar el nombre de un archivo de audio.", "Media", "Implementado"),
    ("RF-017", "Audios", "Eliminar audio", "El sistema debe eliminar el audio, su archivo físico y análisis asociados.", "Alta", "Implementado"),
    ("RF-018", "Audios", "Validación de formato", "El sistema debe validar que los archivos sean de formato permitido (mp3, wav, m4a, ogg, ogm, mp4).", "Alta", "Implementado"),
    ("RF-019", "Audios", "Límite de tamaño", "El sistema debe rechazar archivos mayores a 500 MB.", "Alta", "Implementado"),
    ("RF-020", "Audios", "Detección de duración", "El sistema debe detectar y almacenar la duración del audio (FFprobe/wave).", "Media", "Implementado"),
    ("RF-021", "Audios", "Estadísticas de audios", "El sistema debe proporcionar resumen estadístico: total, transcritos, costo total.", "Media", "Implementado"),
    # Transcription
    ("RF-022", "Transcripción", "Transcripción con Deepgram", "El sistema debe transcribir audios usando Deepgram nova-3 con diarización.", "Alta", "Implementado"),
    ("RF-023", "Transcripción", "Transcripción con WhisperX", "El sistema debe transcribir audios usando WhisperX con alineación temporal.", "Alta", "Implementado"),
    ("RF-024", "Transcripción", "Selección de proveedor", "El usuario debe poder elegir el proveedor de transcripción (Deepgram/WhisperX).", "Alta", "Implementado"),
    ("RF-025", "Transcripción", "Transcripción masiva", "El sistema debe transcribir todos los audios de una campaña en batch.", "Alta", "Implementado"),
    ("RF-026", "Transcripción", "Diarización de hablantes", "La transcripción debe identificar y etiquetar diferentes hablantes.", "Alta", "Implementado"),
    ("RF-027", "Transcripción", "Formato timestamped", "La transcripción debe incluir marcas de tiempo (TIME - TIME | SPEAKER | TEXT).", "Alta", "Implementado"),
    ("RF-028", "Transcripción", "Detener transcripción", "El usuario debe poder detener un proceso de transcripción masiva en curso.", "Media", "Implementado"),
    ("RF-029", "Transcripción", "Monitoreo de progreso", "El sistema debe informar el progreso de la transcripción (completados/total/fallidos).", "Alta", "Implementado"),
    ("RF-030", "Transcripción", "Cálculo de costos", "El sistema debe calcular y almacenar el costo de cada transcripción.", "Media", "Implementado"),
    # Analysis
    ("RF-031", "Análisis", "Análisis con OpenAI", "El sistema debe analizar transcripciones usando OpenAI GPT con prompts personalizados.", "Alta", "Implementado"),
    ("RF-032", "Análisis", "Prompts personalizables", "El usuario debe poder crear, editar y eliminar plantillas de prompts de análisis.", "Alta", "Implementado"),
    ("RF-033", "Análisis", "Formatos de salida", "El usuario debe poder definir formatos de salida con campos y layout personalizado.", "Alta", "Implementado"),
    ("RF-034", "Análisis", "Análisis masivo", "El sistema debe analizar todos los audios transcritos de una campaña.", "Alta", "Implementado"),
    ("RF-035", "Análisis", "Detener análisis", "El usuario debe poder detener un proceso de análisis masivo en curso.", "Media", "Implementado"),
    ("RF-036", "Análisis", "Monitoreo de análisis", "El sistema debe informar el progreso del análisis (completados/total/fallidos).", "Alta", "Implementado"),
    ("RF-037", "Análisis", "Evaluación estructurada", "Cada criterio debe evaluarse como: Cumple, No cumple, No aplica.", "Alta", "Implementado"),
    ("RF-038", "Análisis", "Justificación de evaluación", "Cada evaluación debe incluir justificación detallada y observaciones adicionales.", "Alta", "Implementado"),
    ("RF-039", "Análisis", "Control de tokens/costos", "El sistema debe rastrear tokens consumidos y costos por cada análisis.", "Media", "Implementado"),
    ("RF-040", "Análisis", "Resultados por campaña", "El sistema debe listar resultados de análisis agrupados por campaña.", "Alta", "Implementado"),
    # Pipeline
    ("RF-041", "Pipeline", "Pipeline integrado", "El sistema debe ejecutar secuencialmente: transcripción → análisis en un solo comando.", "Alta", "Implementado"),
    ("RF-042", "Pipeline", "Pipeline asíncrono", "El pipeline debe ejecutarse en background con monitoreo de progreso.", "Alta", "Implementado"),
    # Export
    ("RF-043", "Exportación", "Exportar XLSX", "El sistema debe exportar resultados de análisis en formato Excel.", "Alta", "Implementado"),
    ("RF-044", "Exportación", "Exportar CSV", "El sistema debe exportar resultados de análisis en formato CSV.", "Media", "Pendiente"),
    ("RF-045", "Exportación", "Exportar JSON", "El sistema debe exportar resultados de análisis en formato JSON.", "Media", "Pendiente"),
    ("RF-046", "Exportación", "Incluir transcripción", "La exportación debe permitir incluir o excluir la transcripción.", "Media", "Implementado"),
    ("RF-047", "Exportación", "Metadatos de audio", "La exportación debe extraer metadatos del nombre de archivo (DNI, ejecutivo, fecha).", "Media", "Implementado"),
    # Dashboard & UI
    ("RF-048", "Dashboard", "Estadísticas generales", "El dashboard debe mostrar: total audios, % transcritos, campañas activas, costos.", "Media", "Implementado"),
    ("RF-049", "Dashboard", "Audios recientes", "El dashboard debe listar los audios más recientes.", "Baja", "Implementado"),
    ("RF-050", "Dashboard", "Desglose de costos", "El dashboard debe mostrar costos por campaña (top 8).", "Baja", "Implementado"),
    ("RF-051", "UI", "Tema claro/oscuro", "El sistema debe soportar modo claro y oscuro con persistencia.", "Baja", "Implementado"),
    ("RF-052", "UI", "Navegación con sidebar", "El frontend debe incluir sidebar colapsable con navegación por módulos.", "Media", "Implementado"),
    ("RF-053", "UI", "Indicador de carga", "El frontend debe mostrar indicadores visuales durante operaciones asíncronas.", "Media", "Implementado"),
    # Config suggestions
    ("RF-054", "Configuración", "Sugerencias IA", "El sistema debe sugerir combinaciones de prompt y formato basadas en plantillas existentes.", "Baja", "Implementado"),
    ("RF-055", "Configuración", "Tipos de extracción", "El sistema debe listar tipos de extracción de metadatos disponibles.", "Baja", "Implementado"),
]

create_table(
    doc,
    ["ID", "Módulo", "Requerimiento", "Descripción", "Prioridad", "Estado"],
    reqs_func,
    col_widths=[1.5, 2.2, 2.8, 6, 1.5, 2.2],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 6. REQUERIMIENTOS NO FUNCIONALES
# ══════════════════════════════════════════════════════════════════════════════

add_heading("6. TABLA DE REQUERIMIENTOS NO FUNCIONALES", level=1)

reqs_nofunc = [
    ("RNF-001", "Rendimiento", "El sistema debe responder a peticiones simples en menos de 500ms.", "Alta", "Implementado"),
    ("RNF-002", "Rendimiento", "Las operaciones de carga masiva deben procesarse en background sin bloquear la interfaz.", "Alta", "Implementado"),
    ("RNF-003", "Rendimiento", "El frontend debe implementar lazy loading para reducir el tiempo de carga inicial.", "Alta", "Implementado"),
    ("RNF-004", "Rendimiento", "El frontend debe implementar caché en memoria para reducir peticiones repetitivas.", "Media", "Implementado"),
    ("RNF-005", "Seguridad", "Las contraseñas deben almacenarse hasheadas con bcrypt.", "Alta", "Implementado"),
    ("RNF-006", "Seguridad", "La autenticación debe realizarse mediante tokens JWT con expiración de 60 minutos.", "Alta", "Implementado"),
    ("RNF-007", "Seguridad", "El sistema debe implementar CORS restringido a orígenes confiables.", "Alta", "Implementado"),
    ("RNF-008", "Seguridad", "Cada usuario solo debe ver y gestionar sus propios datos (aislamiento multi-tenant).", "Alta", "Implementado"),
    ("RNF-009", "Seguridad", "Las API keys de servicios externos deben almacenarse en variables de entorno.", "Alta", "Implementado"),
    ("RNF-010", "Seguridad", "El interceptor HTTP debe manejar respuestas 401 redirigiendo al login.", "Alta", "Implementado"),
    ("RNF-011", "Escalabilidad", "El backend debe soportar múltiples API keys de Deepgram con failover automático.", "Media", "Implementado"),
    ("RNF-012", "Escalabilidad", "La arquitectura debe permitir agregar nuevos proveedores de transcripción sin cambios mayores.", "Media", "Implementado"),
    ("RNF-013", "Usabilidad", "La interfaz debe ser responsiva y funcional en resoluciones desktop (1024px+).", "Alta", "Implementado"),
    ("RNF-014", "Usabilidad", "Los formularios deben mostrar mensajes de error en tiempo real.", "Alta", "Implementado"),
    ("RNF-015", "Usabilidad", "La carga de archivos debe soportar drag-and-drop.", "Media", "Implementado"),
    ("RNF-016", "Disponibilidad", "La base de datos debe ejecutarse en contenedor Docker para facilitar despliegue.", "Alta", "Implementado"),
    ("RNF-017", "Mantenibilidad", "El código backend debe seguir el patrón Repository para acceso a datos.", "Alta", "Implementado"),
    ("RNF-018", "Mantenibilidad", "El código frontend debe usar componentes standalone y lazy loading.", "Alta", "Implementado"),
    ("RNF-019", "Mantenibilidad", "Los path aliases (@ prefixed) deben usarse para imports limpios en el frontend.", "Media", "Implementado"),
    ("RNF-020", "Mantenibilidad", "Las variables de entorno deben gestionarse con archivos .env.", "Alta", "Implementado"),
    ("RNF-021", "Portabilidad", "El backend debe ejecutarse en cualquier sistema con Python 3.11+ y PostgreSQL.", "Media", "Implementado"),
    ("RNF-022", "Portabilidad", "El frontend debe compilar como SPA estática desplegable en cualquier servidor web.", "Media", "Implementado"),
    ("RNF-023", "Fiabilidad", "El sistema debe validar formatos y tamaños de archivo antes de procesarlos.", "Alta", "Implementado"),
    ("RNF-024", "Fiabilidad", "Los procesos batch deben poder detenerse manualmente sin corromper datos.", "Alta", "Implementado"),
    ("RNF-025", "Fiabilidad", "Los errores de servicios externos no deben causar caídas del sistema.", "Alta", "Implementado"),
]

create_table(
    doc,
    ["ID", "Categoría", "Descripción", "Prioridad", "Estado"],
    reqs_nofunc,
    col_widths=[1.8, 2.5, 8.5, 1.5, 2.2],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7. ESPECIFICACIÓN DE MÓDULOS
# ══════════════════════════════════════════════════════════════════════════════

add_heading("7. ESPECIFICACIÓN DE MÓDULOS", level=1)

# --- 7.1 Auth ---
add_heading("7.1. Módulo de Autenticación", level=2)
doc.add_paragraph("Responsable de la gestión de identidad y acceso al sistema.")

p = doc.add_paragraph()
p.add_run("Backend:").bold = True
auth_be = [
    ("auth_router.py", "5 endpoints: register, login, me (GET/PUT/DELETE)"),
    ("user_service.py", "UserRepository: CRUD de usuarios"),
    ("auth_deps.py", "Dependencia get_current_user: validación JWT → User"),
    ("security_utils.py", "Hash bcrypt, generación/validación JWT (HS256)"),
]
create_table(doc, ["Archivo", "Descripción"], auth_be, col_widths=[4, 12.5])

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Frontend:").bold = True
auth_fe = [
    ("auth.service.ts", "Servicio: login(), logout(), me(), checkSession(). Signals: authStatus, authMe, token."),
    ("auth.interceptors.ts", "Interceptor HTTP: inyección de Bearer token, manejo 401."),
    ("login-page.ts", "Componente: formulario de login con validación reactiva."),
    ("auth-layout.ts", "Layout wrapper para páginas de autenticación."),
    ("auth-guard.ts", "Guard de ruta: checkSession() → redirige a /auth/login si no autenticado."),
    ("guest-guard.ts", "Guard de ruta: redirige a / si ya está autenticado."),
]
create_table(doc, ["Archivo", "Descripción"], auth_fe, col_widths=[4, 12.5])

# --- 7.2 Campaigns ---
add_heading("7.2. Módulo de Campañas", level=2)
doc.add_paragraph("Gestión del ciclo de vida de campañas de llamadas, incluyendo transcripción masiva, análisis y exportación.")

p = doc.add_paragraph()
p.add_run("Backend:").bold = True
camp_be = [
    ("campaigns_router.py", "18 endpoints: CRUD, stats, transcription, analysis, pipeline, export."),
    ("audio_service.py", "CampaignRepository: CRUD. AudioRepository: consultas filtradas."),
    ("deepgram_service.py", "Transcripción cloud con failover de API keys."),
    ("whisperx_service.py", "Transcripción local con diarización y alineación."),
    ("openai_service.py", "Análisis QA con JSON estructurado."),
]
create_table(doc, ["Archivo", "Descripción"], camp_be, col_widths=[4, 12.5])

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Frontend:").bold = True
camp_fe = [
    ("campaigns.service.ts", "Servicio: load(), create(), update(), delete(), loadTop100WithStats(). Caché paginado."),
    ("campaigns-home.ts", "Página principal: orquesta header, tabla y modales."),
    ("table-campaigns.ts", "Tabla paginada con búsqueda debounced (300ms)."),
    ("add-campaign.ts", "Modal: formulario de creación con validación."),
    ("update-campaign.ts", "Modal: formulario de edición."),
    ("delete-campaign.ts", "Modal: confirmación de eliminación."),
    ("header-campaigns.ts", "Header de página con acciones principales."),
]
create_table(doc, ["Archivo", "Descripción"], camp_fe, col_widths=[4, 12.5])

# --- 7.3 Audios ---
add_heading("7.3. Módulo de Audios", level=2)
doc.add_paragraph("Gestión de archivos de audio: carga, almacenamiento, listado y eliminación.")

p = doc.add_paragraph()
p.add_run("Backend:").bold = True
aud_be = [
    ("audios_router.py", "9 endpoints: CRUD, upload (single/multi), transcribe, stats, delete all by campaign."),
    ("audio_service.py", "AudioRepository: CRUD, consultas con filtros y paginación."),
    ("storage_service.py", "Validación de archivos, almacenamiento en disco, detección de duración."),
    ("ffmpeg_utils.py", "Conversión de formatos de audio. Extracción de duración con ffprobe."),
]
create_table(doc, ["Archivo", "Descripción"], aud_be, col_widths=[4, 12.5])

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Frontend:").bold = True
aud_fe = [
    ("audios.service.ts", "Servicio: load(), uploadAudios(), uploadMultiAudios(), update(), delete(). Invalidación de caché con regex."),
    ("audios-home.ts", "Página principal: orquesta componentes."),
    ("table-audios.ts", "Tabla paginada con filtro por campaña y búsqueda."),
    ("add-audios.ts", "Modal: carga unitaria/masiva con drag-and-drop."),
    ("update-audios.ts", "Modal: edición de nombre de audio."),
    ("delete-audios.ts", "Modal: confirmación de eliminación."),
    ("header-audios.ts", "Header de página."),
]
create_table(doc, ["Archivo", "Descripción"], aud_fe, col_widths=[4, 12.5])

# --- 7.4 Analysis ---
add_heading("7.4. Módulo de Análisis", level=2)
doc.add_paragraph(
    "Centro de análisis QA: configuración de pipeline, gestión de prompts y formatos, "
    "ejecución de análisis con IA y visualización de resultados."
)

p = doc.add_paragraph()
p.add_run("Backend:").bold = True
ana_be = [
    ("analyses_router.py", "3 endpoints: listar, crear manual, eliminar análisis."),
    ("analysis_service.py", "AnalysisRepository: create_batch() para inserción masiva de resultados."),
    ("openai_service.py", "Servicio central: prompt → OpenAI GPT → JSON estructurado → normalización → almacenamiento."),
    ("cost_utils.py", "CostTracker: cálculo de costos por tokens (in/out) con precios por modelo."),
]
create_table(doc, ["Archivo", "Descripción"], ana_be, col_widths=[4, 12.5])

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Frontend:").bold = True
ana_fe = [
    ("analysis.service.ts", "Servicio central: CRUD prompts/formatos, start/stop/status/results de análisis. Pipeline."),
    ("analysis-home.ts", "Página principal: orquesta 5 sub-componentes."),
    ("process-analysis.ts", "Configuración del pipeline: selección de campaña, proveedor, modo (transcribir/analizar/ambos)."),
    ("prompt-manager.ts", "CRUD de plantillas de prompts."),
    ("format-manager.ts", "CRUD de formatos de salida con campos JSON y layout configurable."),
    ("results-analysis-table.ts", "Tabla de resultados con paginación y filtros. Visualización de transcripciones."),
    ("header-analysis.ts", "Header de página."),
]
create_table(doc, ["Archivo", "Descripción"], ana_fe, col_widths=[4, 12.5])

# --- 7.5 Config ---
add_heading("7.5. Módulo de Configuración", level=2)
doc.add_paragraph("Gestión de plantillas de prompts y formatos de salida reutilizables.")

config_be = [
    ("configs_router.py", "10 endpoints: CRUD prompts, CRUD formatos, sugerencias, tipos de metadatos."),
    ("config_service.py", "PromptTemplateRepository, OutputFormatRepository: CRUD con filtrado por usuario."),
]
create_table(doc, ["Archivo", "Descripción"], config_be, col_widths=[4, 12.5])

# --- 7.6 Dashboard ---
add_heading("7.6. Módulo Dashboard", level=2)
doc.add_paragraph("Panel de control con métricas generales y resumen de actividad.")

dash = [
    ("dashboard.ts (frontend)", "Componente standalone que consume múltiples endpoints para renderizar estadísticas."),
    ("GET /audios/stats/summary", "Total audios, transcritos, costo total."),
    ("GET /campaigns/with-stats", "Campañas con estadísticas de audios y costos."),
]
create_table(doc, ["Componente / Endpoint", "Descripción"], dash, col_widths=[5, 11.5])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 8. ESPECIFICACIÓN DE ENDPOINTS API
# ══════════════════════════════════════════════════════════════════════════════

add_heading("8. ESPECIFICACIÓN DE ENDPOINTS API", level=1)

doc.add_paragraph("Base URL: http://localhost:8000 (desarrollo)")
doc.add_paragraph("Autenticación: Bearer JWT en header Authorization (excepto register/login)")

# 8.1 Auth
add_heading("8.1. Auth API (/auth)", level=2)
auth_eps = [
    ("POST", "/auth/register", "Registrar usuario", "RegisterRequest (username, password)", "UserResponse (201)"),
    ("POST", "/auth/login", "Iniciar sesión", "LoginRequest (username, password)", "TokenResponse (access_token)"),
    ("GET", "/auth/me", "Obtener perfil", "-", "UserResponse"),
    ("PUT", "/auth/me", "Actualizar perfil", "UpdateUserRequest (username?, password?)", "UserResponse"),
    ("DELETE", "/auth/me", "Eliminar cuenta", "-", "204 No Content"),
]
create_table(doc, ["Método", "Ruta", "Descripción", "Request", "Response"], auth_eps, col_widths=[1.5, 3, 3, 4.5, 4.5])

# 8.2 Campaigns
add_heading("8.2. Campaigns API (/campaigns)", level=2)
camp_eps = [
    ("POST", "/campaigns/", "Crear campaña", "CampaignCreate", "CampaignResponse (201)"),
    ("GET", "/campaigns/", "Listar campañas", "Query: page, page_size, search", "CampaignPage"),
    ("GET", "/campaigns/with-stats", "Campañas con stats", "Query: page, page_size", "Custom object"),
    ("GET", "/campaigns/{id}", "Detalle campaña", "-", "CampaignResponse"),
    ("PUT", "/campaigns/{id}", "Actualizar campaña", "CampaignUpdate", "CampaignResponse"),
    ("DELETE", "/campaigns/{id}", "Eliminar campaña", "-", "MessageResponse (204)"),
    ("GET", "/campaigns/{id}/stats", "Estadísticas", "-", "Stats object"),
    ("POST", "/campaigns/{id}/transcribe-all", "Transcripción masiva", "provider (deepgram/whisperx)", "Status message"),
    ("POST", "/campaigns/{id}/transcribe-stop", "Detener transcripción", "-", "Status message"),
    ("GET", "/campaigns/{id}/transcribe-status", "Progreso transcripción", "-", "AsyncStatus"),
    ("POST", "/campaigns/{id}/analyze-all", "Análisis masivo sync", "CampaignAnalysisRequest", "Results"),
    ("POST", "/campaigns/{id}/analyze-all-async", "Análisis masivo async", "AsyncAnalysisRequest", "Queue response"),
    ("POST", "/campaigns/{id}/pipeline-async", "Pipeline completo", "PipelineRequest", "Queue response"),
    ("POST", "/campaigns/{id}/analyze-stop", "Detener análisis", "-", "Status message"),
    ("GET", "/campaigns/{id}/analyze-status", "Progreso análisis", "-", "AsyncStatus"),
    ("GET", "/campaigns/{id}/analysis-results", "Resultados análisis", "-", "Results array"),
    ("GET", "/campaigns/{id}/transcriptions", "Transcripciones", "-", "Transcriptions array"),
    ("GET", "/campaigns/{id}/analysis-export", "Exportar XLSX", "-", "File stream (.xlsx)"),
]
create_table(doc, ["Método", "Ruta", "Descripción", "Request", "Response"], camp_eps, col_widths=[1.3, 4.5, 2.5, 4, 4])

# 8.3 Audios
add_heading("8.3. Audios API (/audios)", level=2)
aud_eps = [
    ("GET", "/audios/", "Listar audios", "Query: page, page_size, campaign_id, search", "AudioPage"),
    ("GET", "/audios/stats/summary", "Estadísticas", "-", "Stats object"),
    ("GET", "/audios/{id}", "Detalle audio", "-", "AudioResponse"),
    ("POST", "/audios/upload", "Subir audio", "Form: file, campaign_id", "AudioResponse (201)"),
    ("POST", "/audios/upload-multiple", "Subir múltiples", "Form: files[], campaign_id", "AudioResponse[]"),
    ("PUT", "/audios/{id}", "Actualizar audio", "AudioUpdateRequest", "AudioResponse"),
    ("DELETE", "/audios/{id}", "Eliminar audio", "-", "MessageResponse (204)"),
    ("POST", "/audios/{id}/transcribe", "Transcribir audio", "provider", "Status message"),
    ("DELETE", "/audios/campaign/{id}/all", "Eliminar todos de campaña", "-", "MessageResponse"),
]
create_table(doc, ["Método", "Ruta", "Descripción", "Request", "Response"], aud_eps, col_widths=[1.3, 4, 2.5, 5, 3.5])

# 8.4 Analysis
add_heading("8.4. Analyses API (/analyses)", level=2)
ana_eps = [
    ("GET", "/analyses/", "Listar análisis", "Query: audio_id?, page, page_size", "AnalysisPage"),
    ("POST", "/analyses/", "Crear análisis manual", "AnalysisCreate", "AnalysisResponse (201)"),
    ("DELETE", "/analyses/{id}", "Eliminar análisis", "-", "MessageResponse (204)"),
]
create_table(doc, ["Método", "Ruta", "Descripción", "Request", "Response"], ana_eps, col_widths=[1.5, 3, 3.5, 4.5, 4])

# 8.5 Config
add_heading("8.5. Config API (/analysis-configs)", level=2)
conf_eps = [
    ("GET", "/analysis-configs/prompts", "Listar prompts", "-", "PromptTemplate[]"),
    ("POST", "/analysis-configs/prompts", "Crear prompt", "PromptTemplateCreate", "PromptTemplateResponse"),
    ("PUT", "/analysis-configs/prompts/{id}", "Actualizar prompt", "PromptTemplateUpdate", "PromptTemplateResponse"),
    ("DELETE", "/analysis-configs/prompts/{id}", "Eliminar prompt", "-", "MessageResponse"),
    ("GET", "/analysis-configs/formats", "Listar formatos", "-", "OutputFormat[]"),
    ("POST", "/analysis-configs/formats", "Crear formato", "OutputFormatCreate", "OutputFormatResponse"),
    ("PUT", "/analysis-configs/formats/{id}", "Actualizar formato", "OutputFormatUpdate", "OutputFormatResponse"),
    ("DELETE", "/analysis-configs/formats/{id}", "Eliminar formato", "-", "MessageResponse"),
    ("GET", "/analysis-configs/prompt-format-suggestions", "Sugerencias IA", "-", "Suggestions[]"),
    ("GET", "/analysis-configs/metadata-extraction-types", "Tipos extracción", "-", "MetadataTypes[]"),
]
create_table(doc, ["Método", "Ruta", "Descripción", "Request", "Response"], conf_eps, col_widths=[1.3, 5.2, 2.5, 3.5, 4])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 9. ESTRUCTURA DEL PROYECTO
# ══════════════════════════════════════════════════════════════════════════════

add_heading("9. ESTRUCTURA DEL PROYECTO", level=1)

add_heading("9.1. Estructura del Backend", level=2)
backend_struct = """backend-transcription-analysis/
├── main.py                         # Inicialización FastAPI, CORS, registro de routers
├── create_tables.py                # Script de creación de tablas
├── seed_analysis_prompts.py        # Seed de prompts predefinidos
├── compose.yaml                    # Docker Compose (PostgreSQL)
├── requirements.txt                # Dependencias Python
├── .env                            # Variables de entorno
├── .env.template                   # Template de variables
├── uploads/                        # Directorio de archivos de audio
└── app/
    ├── components/
    │   ├── connection.py           # SQLAlchemy engine, SessionLocal, get_db()
    │   ├── models.py              # 6 modelos ORM
    │   └── schemas.py             # 20+ schemas Pydantic
    ├── routers/
    │   ├── auth_router.py          # Auth endpoints
    │   ├── campaigns_router.py     # Campaign endpoints
    │   ├── audios_router.py        # Audio endpoints
    │   ├── analyses_router.py      # Analysis endpoints
    │   └── configs_router.py       # Config endpoints
    ├── services/
    │   ├── user_service.py         # UserRepository
    │   ├── audio_service.py        # AudioRepository, CampaignRepository
    │   ├── analysis_service.py     # AnalysisRepository
    │   ├── deepgram_service.py     # Deepgram transcription
    │   ├── whisperx_service.py     # WhisperX transcription
    │   ├── openai_service.py       # OpenAI analysis
    │   ├── config_service.py       # Prompt/Format repositories
    │   └── storage_service.py      # File storage
    ├── dependencies/
    │   └── auth_deps.py            # JWT validation dependency
    └── utils/
        ├── security_utils.py       # Hashing, JWT
        ├── extra_utils.py          # Audio, device, datetime utilities
        ├── file_utils.py           # File operations
        ├── ffmpeg_utils.py         # Audio conversion
        └── cost_utils.py           # Cost tracking"""

p = doc.add_paragraph()
run = p.add_run(backend_struct)
run.font.name = "Consolas"
run.font.size = Pt(8)

add_heading("9.2. Estructura del Frontend", level=2)
frontend_struct = """frontend-transcription-analysis/
├── angular.json                    # Configuración Angular CLI
├── package.json                    # Dependencias npm
├── tsconfig.json                   # TypeScript config (path aliases)
├── public/                         # Assets estáticos
├── scripts/
│   └── set-envs.js                # Script de configuración de entornos
└── src/
    ├── main.ts                     # Bootstrap de la aplicación
    ├── index.html                  # HTML raíz
    ├── styles.css                  # Estilos globales (Tailwind)
    ├── environments/
    │   ├── environment.ts          # Config producción
    │   └── environment.development.ts  # Config desarrollo
    └── app/
        ├── app.ts                  # Componente raíz (RouterOutlet)
        ├── app.config.ts           # Providers: HttpClient, Router, Interceptors
        ├── app.routes.ts           # Rutas principales con lazy loading
        ├── api/                    # Interfaces TypeScript
        │   ├── auth.interface.ts
        │   ├── campaigns.interface.ts
        │   ├── audios.interface.ts
        │   ├── analysis.interface.ts
        │   └── page.interface.ts
        ├── core/                   # Servicios singleton
        │   ├── cache/cache.service.ts
        │   ├── config/api-config.service.ts
        │   ├── storage/token-storage.service.ts
        │   └── theme/theme.service.ts
        ├── feactures/              # Módulos de negocio
        │   ├── auth/               # Autenticación
        │   ├── campaigns/          # Campañas
        │   ├── audios/             # Audios
        │   ├── analysis/           # Análisis
        │   └── shell/              # Layout (home, dashboard)
        ├── guards/                 # Route guards
        │   ├── auth-guard.ts
        │   └── guest-guard.ts
        └── shared/                 # Componentes reutilizables
            ├── navbar/
            ├── sidebar/
            ├── footer/
            ├── loading/
            ├── theme-buttons/
            └── utils/form.utils.ts"""

p = doc.add_paragraph()
run = p.add_run(frontend_struct)
run.font.name = "Consolas"
run.font.size = Pt(8)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 10. FLUJOS DE PROCESO PRINCIPALES
# ══════════════════════════════════════════════════════════════════════════════

add_heading("10. FLUJOS DE PROCESO PRINCIPALES", level=1)

add_heading("10.1. Pipeline de Transcripción", level=2)
trans_steps = [
    ("1", "Selección", "El usuario selecciona una campaña y el proveedor de transcripción (Deepgram o WhisperX)."),
    ("2", "Inicio", "Frontend envía POST /campaigns/{id}/transcribe-all con el proveedor seleccionado."),
    ("3", "Background", "Backend crea una tarea en background. Registra un job tracker con (total, completed, failed)."),
    ("4", "Procesamiento", "Para cada audio sin transcripción: carga archivo → envía al proveedor → recibe texto con timestamps y speakers."),
    ("5", "Almacenamiento", "Guarda la transcripción en Audio.transcription, el costo en Audio.cost y la duración en Audio.minutes."),
    ("6", "Monitoreo", "Frontend hace polling cada 3 segundos a GET /campaigns/{id}/transcribe-status."),
    ("7", "Finalización", "Job tracker marca como completado. Frontend muestra resultados. Caché invalidado."),
    ("8", "Cancelación", "Si el usuario solicita detener: POST /campaigns/{id}/transcribe-stop marca cancelled=True."),
]
create_table(doc, ["Paso", "Fase", "Descripción"], trans_steps, col_widths=[1, 2.5, 13])

add_heading("10.2. Pipeline de Análisis con IA", level=2)
analysis_steps = [
    ("1", "Configuración", "El usuario selecciona: campaña, prompt de análisis, formato de salida, y opcionalmente tipo de extracción de metadatos."),
    ("2", "Inicio", "Frontend envía POST /campaigns/{id}/analyze-all-async con prompt_id, format_id, y opciones."),
    ("3", "Background", "Backend crea tarea en background. Obtiene todos los audios transcritos de la campaña."),
    ("4", "Procesamiento", "Para cada audio: elimina análisis previos → construye prompt con transcripción → envía a OpenAI GPT."),
    ("5", "Respuesta IA", "OpenAI responde con JSON estructurado: array de {criterio, evaluacion, justificacion, obs_adicional}."),
    ("6", "Normalización", "Backend normaliza campos (mapea sinónimos), valida evaluación (Cumple/No cumple/No aplica)."),
    ("7", "Almacenamiento", "Crea batch de AudioAnalysis: cada criterio = 1 registro. Almacena tokens, costo, result_json."),
    ("8", "Monitoreo", "Frontend hace polling cada 3 segundos a GET /campaigns/{id}/analyze-status."),
    ("9", "Resultados", "Frontend carga GET /campaigns/{id}/analysis-results y renderiza tabla de resultados."),
]
create_table(doc, ["Paso", "Fase", "Descripción"], analysis_steps, col_widths=[1, 2.5, 13])

add_heading("10.3. Flujo de Autenticación", level=2)
auth_steps = [
    ("1", "Login", "Usuario ingresa username/password en formulario con validación reactiva."),
    ("2", "Request", "Frontend envía POST /auth/login con credenciales."),
    ("3", "Validación", "Backend verifica username existe → compara password con hash bcrypt almacenado."),
    ("4", "Token", "Si válido: genera JWT (HS256) con sub=user_id, exp=60min. Retorna access_token."),
    ("5", "Almacenamiento", "Frontend guarda token en sessionStorage vía TokenStorageService."),
    ("6", "Interceptor", "En cada petición HTTP, authInterceptor inyecta header Authorization: Bearer {token}."),
    ("7", "Guard", "Al navegar, AuthGuard ejecuta checkSession(): llama GET /auth/me para verificar token."),
    ("8", "401 Handler", "Si cualquier respuesta es 401: interceptor limpia sesión y redirige a /auth/login."),
    ("9", "Logout", "AuthService.logout(): limpia token, estado y caché. Navega a /auth/login."),
]
create_table(doc, ["Paso", "Fase", "Descripción"], auth_steps, col_widths=[1, 2.5, 13])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 11. DISEÑO DE BASE DE DATOS
# ══════════════════════════════════════════════════════════════════════════════

add_heading("11. DISEÑO DE LA BASE DE DATOS", level=1)

doc.add_paragraph(
    "La base de datos está implementada en PostgreSQL 16 y utiliza 6 tablas principales. "
    "Todas las tablas usan UUID como clave primaria generado automáticamente y timestamps "
    "de creación/actualización automáticos."
)

add_heading("Tabla: tbl_users", level=3)
users_cols = [
    ("id", "UUID", "PK", "Identificador único generado automáticamente"),
    ("username", "VARCHAR(50)", "UNIQUE, NOT NULL", "Nombre de usuario"),
    ("password", "TEXT", "NOT NULL", "Contraseña hasheada con bcrypt"),
    ("register_date", "TIMESTAMP", "DEFAULT now()", "Fecha de registro"),
    ("updated_date", "TIMESTAMP", "ON UPDATE now()", "Fecha de última actualización"),
]
create_table(doc, ["Columna", "Tipo", "Restricciones", "Descripción"], users_cols, col_widths=[3, 3, 4, 6.5])

add_heading("Tabla: tbl_campaigns", level=3)
camp_cols = [
    ("id", "UUID", "PK", "Identificador único"),
    ("user_id", "UUID", "FK → tbl_users.id CASCADE", "Propietario de la campaña"),
    ("campaign_name", "VARCHAR(255)", "NOT NULL", "Nombre de la campaña"),
    ("description", "TEXT", "NULLABLE", "Descripción opcional"),
    ("register_date", "TIMESTAMP", "DEFAULT now()", "Fecha de creación"),
    ("updated_date", "TIMESTAMP", "ON UPDATE now()", "Última actualización"),
]
create_table(doc, ["Columna", "Tipo", "Restricciones", "Descripción"], camp_cols, col_widths=[3, 3, 4, 6.5])

add_heading("Tabla: tbl_audios", level=3)
aud_cols = [
    ("id", "UUID", "PK", "Identificador único"),
    ("campaign_id", "UUID", "FK → tbl_campaigns.id CASCADE", "Campaña asociada"),
    ("audio_name", "VARCHAR(255)", "NOT NULL", "Nombre del archivo de audio"),
    ("transcription", "TEXT", "NULLABLE", "Transcripción timestamped con speakers"),
    ("cost", "NUMERIC(10,4)", "DEFAULT 0", "Costo de transcripción en USD"),
    ("minutes", "NUMERIC(10,2)", "DEFAULT 0", "Duración del audio en minutos"),
    ("register_date", "TIMESTAMP", "DEFAULT now()", "Fecha de carga"),
    ("updated_date", "TIMESTAMP", "ON UPDATE now()", "Última actualización"),
]
create_table(doc, ["Columna", "Tipo", "Restricciones", "Descripción"], aud_cols, col_widths=[3, 3, 4, 6.5])

add_heading("Tabla: tbl_audios_analysis", level=3)
ana_cols = [
    ("id", "UUID", "PK", "Identificador único"),
    ("audio_id", "UUID", "FK → tbl_audios.id CASCADE", "Audio analizado"),
    ("criterio", "TEXT", "NOT NULL", "Nombre del criterio evaluado"),
    ("evaluacion", "VARCHAR(50)", "NOT NULL", "Cumple / No cumple / No aplica"),
    ("justificacion", "TEXT", "NOT NULL", "Justificación de la evaluación"),
    ("obs_adicional", "TEXT", "NULLABLE", "Observaciones adicionales"),
    ("in_token", "INTEGER", "NULLABLE", "Tokens de entrada consumidos"),
    ("out_token", "INTEGER", "NULLABLE", "Tokens de salida consumidos"),
    ("cost", "NUMERIC(10,4)", "DEFAULT 0", "Costo del análisis en USD"),
    ("result_json", "TEXT", "NULLABLE", "JSON completo de la respuesta OpenAI"),
    ("register_date", "TIMESTAMP", "DEFAULT now()", "Fecha del análisis"),
    ("updated_date", "TIMESTAMP", "ON UPDATE now()", "Última actualización"),
]
create_table(doc, ["Columna", "Tipo", "Restricciones", "Descripción"], ana_cols, col_widths=[3, 2.8, 4, 6.7])

add_heading("Tabla: tbl_prompt_templates", level=3)
prompt_cols = [
    ("id", "UUID", "PK", "Identificador único"),
    ("user_id", "UUID", "FK → tbl_users.id CASCADE", "Propietario"),
    ("name", "VARCHAR(120)", "NOT NULL", "Nombre de la plantilla"),
    ("prompt_text", "TEXT", "NOT NULL", "Texto del prompt de análisis"),
    ("is_active", "BOOLEAN", "DEFAULT TRUE", "Estado activo/inactivo"),
    ("register_date", "TIMESTAMP", "DEFAULT now()", "Fecha de creación"),
    ("updated_date", "TIMESTAMP", "ON UPDATE now()", "Última actualización"),
]
create_table(doc, ["Columna", "Tipo", "Restricciones", "Descripción"], prompt_cols, col_widths=[3, 3, 4, 6.5])

add_heading("Tabla: tbl_output_formats", level=3)
format_cols = [
    ("id", "UUID", "PK", "Identificador único"),
    ("user_id", "UUID", "FK → tbl_users.id CASCADE", "Propietario"),
    ("name", "VARCHAR(120)", "NOT NULL", "Nombre del formato"),
    ("fields_json", "TEXT", "NOT NULL", "JSON con lista de campos y layout"),
    ("description", "TEXT", "NULLABLE", "Descripción del formato"),
    ("is_active", "BOOLEAN", "DEFAULT TRUE", "Estado activo/inactivo"),
    ("register_date", "TIMESTAMP", "DEFAULT now()", "Fecha de creación"),
    ("updated_date", "TIMESTAMP", "ON UPDATE now()", "Última actualización"),
]
create_table(doc, ["Columna", "Tipo", "Restricciones", "Descripción"], format_cols, col_widths=[3, 3, 4, 6.5])

add_heading("Relaciones entre Tablas", level=3)
relations = [
    ("tbl_users → tbl_campaigns", "1:N", "Un usuario tiene muchas campañas. CASCADE DELETE."),
    ("tbl_campaigns → tbl_audios", "1:N", "Una campaña tiene muchos audios. CASCADE DELETE."),
    ("tbl_audios → tbl_audios_analysis", "1:N", "Un audio tiene muchos análisis. CASCADE DELETE."),
    ("tbl_users → tbl_prompt_templates", "1:N", "Un usuario tiene muchas plantillas de prompt. CASCADE DELETE."),
    ("tbl_users → tbl_output_formats", "1:N", "Un usuario tiene muchos formatos de salida. CASCADE DELETE."),
]
create_table(doc, ["Relación", "Tipo", "Descripción"], relations, col_widths=[5, 1.5, 10])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 12. GESTIÓN DE CONFIGURACIÓN
# ══════════════════════════════════════════════════════════════════════════════

add_heading("12. GESTIÓN DE CONFIGURACIÓN Y ENTORNO", level=1)

add_heading("12.1. Variables de Entorno del Backend", level=2)
doc.add_paragraph("El backend utiliza un archivo .env para toda la configuración sensible y parametrizable:")

env_vars = [
    ("API_TITLE", "API", "Nombre de la aplicación en la documentación Swagger"),
    ("API_VERSION", "API", "Versión de la API"),
    ("CORS_ORIGINS", "Seguridad", "Orígenes permitidos para CORS (separados por coma)"),
    ("DATABASE_URL", "Base de Datos", "Connection string de PostgreSQL"),
    ("JWT_SECRET", "Autenticación", "Clave secreta para firmar tokens JWT"),
    ("JWT_ALGORITHM", "Autenticación", "Algoritmo de firma (HS256)"),
    ("JWT_EXPIRE_MINUTES", "Autenticación", "Tiempo de expiración del token en minutos"),
    ("DEEPGRAM_KEYS", "Transcripción", "API keys de Deepgram (separadas por coma)"),
    ("DEEPGRAM_MODEL", "Transcripción", "Modelo de Deepgram (nova-3)"),
    ("DEEPGRAM_LANGUAGE", "Transcripción", "Idioma de transcripción (es)"),
    ("WHISPERX_MODEL_SIZE", "Transcripción", "Tamaño del modelo WhisperX (medium)"),
    ("WHISPERX_DEVICE", "Transcripción", "Dispositivo de cómputo (auto/cpu/cuda)"),
    ("HUGGINGFACE_TOKEN", "Diarización", "Token de acceso para modelos HuggingFace"),
    ("OPENAI_API_KEY", "Análisis", "API key de OpenAI"),
    ("OPENAI_MODEL", "Análisis", "Modelo a utilizar (gpt-4.1-mini)"),
    ("OPENAI_MAX_TOKENS", "Análisis", "Máximo de tokens por respuesta"),
    ("OPENAI_TEMPERATURE", "Análisis", "Temperatura del modelo (0.0-1.0)"),
    ("UPLOAD_DIR", "Archivos", "Directorio de almacenamiento de audios"),
    ("MAX_FILE_SIZE", "Archivos", "Tamaño máximo por archivo (bytes)"),
    ("FFMPEG_TIMEOUT", "Audio", "Timeout para operaciones de FFmpeg"),
]
create_table(doc, ["Variable", "Grupo", "Descripción"], env_vars, col_widths=[4, 2.5, 10])

add_heading("12.2. Configuración del Frontend", level=2)
doc.add_paragraph("El frontend gestiona sus entornos mediante archivos en src/environments/:")

fe_env = [
    ("environment.ts", "Producción", "apiUrl apuntando al servidor de producción"),
    ("environment.development.ts", "Desarrollo", "apiUrl: http://localhost:8000 (backend local)"),
    ("set-envs.js", "Script", "Genera archivos de entorno dinámicamente en CI/CD"),
]
create_table(doc, ["Archivo", "Entorno", "Descripción"], fe_env, col_widths=[4.5, 2.5, 9.5])

add_heading("12.3. Path Aliases (Frontend)", level=2)
doc.add_paragraph("El frontend utiliza path aliases configurados en tsconfig.json para imports limpios:")

aliases = [
    ("@shared/*", "src/app/shared/*", "Componentes compartidos"),
    ("@core/*", "src/app/core/*", "Servicios core (singleton)"),
    ("@feactures/*", "src/app/feactures/*", "Módulos de negocio"),
    ("@api/*", "src/app/api/*", "Interfaces de API"),
    ("@guards/*", "src/app/guards/*", "Guards de ruta"),
    ("@environments/*", "src/environments/*", "Variables de entorno"),
]
create_table(doc, ["Alias", "Ruta Real", "Uso"], aliases, col_widths=[3.5, 5, 8])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 13. SEGURIDAD DEL SISTEMA
# ══════════════════════════════════════════════════════════════════════════════

add_heading("13. SEGURIDAD DEL SISTEMA", level=1)

security = [
    ("Hashing de Contraseñas", "bcrypt via passlib", "Las contraseñas nunca se almacenan en texto plano. Se usa bcrypt con salt automático."),
    ("Autenticación JWT", "python-jose (HS256)", "Tokens con expiración configurable (60 min). Claim 'sub' contiene UUID del usuario."),
    ("Interceptor HTTP", "Angular HttpInterceptor", "Inyecta Bearer token en cada request. Maneja 401 con logout automático."),
    ("CORS", "FastAPI CORSMiddleware", "Orígenes restringidos a dominios configurados. Credentials habilitado."),
    ("Aislamiento de datos", "Query-level filtering", "Cada consulta filtra por user_id del token. Un usuario nunca accede a datos de otro."),
    ("Validación de entrada", "Pydantic schemas", "Toda entrada se valida con tipos estrictos, longitudes y formatos antes de procesarse."),
    ("Validación de archivos", "Extensión + tamaño", "Solo se aceptan formatos de audio válidos y archivos menores a 500 MB."),
    ("API Keys en .env", "Variables de entorno", "Ninguna API key hardcodeada. Todas en archivo .env excluido de Git."),
    ("Guards de ruta", "Angular Guards", "AuthGuard protege rutas autenticadas. GuestGuard protege rutas de invitados."),
    ("Token en sessionStorage", "Frontend", "El token se almacena en sessionStorage (se limpia al cerrar pestaña)."),
]
create_table(doc, ["Control", "Implementación", "Descripción"], security, col_widths=[3.5, 3.5, 9.5])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 14. EXPORTACIÓN DE DATOS
# ══════════════════════════════════════════════════════════════════════════════

add_heading("14. EXPORTACIÓN DE DATOS", level=1)

doc.add_paragraph(
    "El sistema ofrece tres formatos de exportación para los resultados de análisis de una campaña. "
    "La exportación incluye datos de análisis, metadatos extraídos del nombre del archivo de audio, "
    "y opcionalmente las transcripciones."
)

export_formats = [
    ("XLSX (Excel)", "GET /campaigns/{id}/analysis-export", "Formato ancho: 1 fila por audio, columnas dinámicas según criterios. Incluye metadatos y observaciones agrupadas."),
    ("CSV", "Pendiente", "Pendiente de implementación. Se planea como extensión del endpoint de exportación XLSX."),
    ("JSON", "Pendiente", "Pendiente de implementación. Se planea como endpoint adicional para consumo programático."),
]
create_table(doc, ["Formato", "Endpoint", "Descripción"], export_formats, col_widths=[2.5, 5, 9])

doc.add_paragraph()
doc.add_paragraph("Estructura de columnas del formato ancho (XLSX/CSV):")
export_cols = [
    ("AUDIO_ID", "UUID del audio"),
    ("AUDIO_NOMBRE", "Nombre del archivo de audio"),
    ("CANTIDAD_CRITERIOS", "Número total de criterios evaluados"),
    ("Metadatos (DNI_EVALUADOR, EJECUTIVO, etc.)", "Extraídos del nombre del archivo según tipo de extracción configurado"),
    ("[1], [2], [3]...", "Columnas dinámicas con evaluación y justificación de cada criterio"),
    ("OBSERVACIONES_X_AL_Y", "Observaciones agrupadas de criterios que no cumplen"),
    ("TRANSCRIPCION_LLAMADA_N", "Transcripción dividida en chunks (opcional)"),
]
create_table(doc, ["Columna", "Descripción"], export_cols, col_widths=[6, 10.5])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 15. GESTIÓN DEL PROYECTO
# ══════════════════════════════════════════════════════════════════════════════

add_heading("15. GESTIÓN DEL PROYECTO", level=1)

add_heading("15.1. Metodología de Desarrollo", level=2)
doc.add_paragraph(
    "El proyecto adopta una metodología de desarrollo ágil iterativa e incremental, combinando "
    "principios de Scrum y Kanban adaptados al tamaño del equipo. El desarrollo se organiza en "
    "sprints cortos (1-2 semanas) con entregas incrementales de funcionalidad."
)

doc.add_paragraph("Principios metodológicos aplicados:")
principles = [
    "Desarrollo orientado a features: cada módulo se desarrolla de forma independiente (backend → frontend).",
    "API-First: los endpoints se diseñan y documentan antes de la implementación del frontend.",
    "Integración continua: cada feature se integra al completarse, no al final del sprint.",
    "Testing incremental: cada endpoint se prueba con la documentación Swagger antes de la integración con el frontend.",
    "Refactoring continuo: mejoras de código se realizan como parte natural del desarrollo.",
]
for p_text in principles:
    doc.add_paragraph(p_text, style="List Bullet")

add_heading("Fases del Desarrollo", level=3)
phases = [
    ("Fase 1 - Fundación", "Configuración del proyecto, stack tecnológico, autenticación, base de datos.", "Completada"),
    ("Fase 2 - CRUD Core", "Módulos de Campañas y Audios: CRUD completo, carga de archivos, paginación.", "Completada"),
    ("Fase 3 - Transcripción", "Integración Deepgram y WhisperX. Transcripción unitaria y masiva.", "Completada"),
    ("Fase 4 - Análisis IA", "Integración OpenAI. Prompts y formatos configurables. Análisis masivo.", "Completada"),
    ("Fase 5 - Pipeline", "Pipeline integrado (transcripción + análisis). Exportación de datos.", "Completada"),
    ("Fase 6 - Dashboard", "Panel de estadísticas. Métricas de costos y progreso.", "Completada"),
    ("Fase 7 - Refinamiento", "Optimización de UX, temas visuales, caché, validaciones.", "Completada"),
]
create_table(doc, ["Fase", "Alcance", "Estado"], phases, col_widths=[3.5, 10, 3])

add_heading("15.2. Control de Versiones", level=2)
doc.add_paragraph("El proyecto utiliza Git como sistema de control de versiones con la siguiente estrategia:")

git_rules = [
    ("Repositorios", "Dos repositorios separados: backend-transcription-analysis y frontend-transcription-analysis."),
    ("Rama principal", "main - Código estable y funcional."),
    ("Ramas feature", "feature/{nombre} - Para desarrollo de nuevas funcionalidades."),
    ("Ramas fix", "fix/{nombre} - Para correcciones de errores."),
    ("Commits", "Mensajes descriptivos en español o inglés. Prefijos: feat:, fix:, refactor:, docs:."),
]
create_table(doc, ["Aspecto", "Convención"], git_rules, col_widths=[3.5, 13])

add_heading("15.3. Convenciones de Código", level=2)

add_heading("Backend (Python)", level=3)
py_conventions = [
    ("Naming", "snake_case para funciones y variables. PascalCase para clases."),
    ("Tipado", "Type hints en parámetros y retornos de funciones."),
    ("Docstrings", "Documentación en funciones públicas de servicios."),
    ("Imports", "Ordenados: stdlib → third-party → local."),
    ("Async", "Funciones async def para endpoints. Sync para lógica de BD."),
    ("Schemas", "Pydantic para toda validación de entrada/salida."),
]
create_table(doc, ["Aspecto", "Convención"], py_conventions, col_widths=[3, 13.5])

add_heading("Frontend (TypeScript/Angular)", level=3)
ts_conventions = [
    ("Componentes", "Standalone (sin NgModules). Archivos separados: .ts, .html, .css."),
    ("Naming", "camelCase para propiedades/métodos. PascalCase para clases/interfaces."),
    ("Estado", "Angular Signals para estado local. RxJS para flujos asíncronos."),
    ("Servicios", "Injectable providedIn: 'root' para singleton."),
    ("Imports", "Path aliases @shared, @core, @feactures, @api, @guards."),
    ("Formularios", "Reactive Forms con FormGroup/FormControl."),
    ("Formato", "Prettier: 100 chars, single quotes, trailing comma. Angular HTML parser."),
]
create_table(doc, ["Aspecto", "Convención"], ts_conventions, col_widths=[3, 13.5])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 16. PLAN DE PRUEBAS
# ══════════════════════════════════════════════════════════════════════════════

add_heading("16. PLAN DE PRUEBAS", level=1)

add_heading("16.1. Estrategia General", level=2)
doc.add_paragraph(
    "El plan de pruebas contempla múltiples niveles de verificación para asegurar "
    "la calidad y estabilidad del sistema."
)

test_types = [
    ("Pruebas Unitarias", "Frontend", "Vitest + JSDOM", "Componentes individuales, servicios, utilidades."),
    ("Pruebas de API", "Backend", "Swagger UI / Postman", "Cada endpoint verificado con datos válidos e inválidos."),
    ("Pruebas de Integración", "Full Stack", "Manual", "Flujos completos: login → crear campaña → subir audio → transcribir → analizar → exportar."),
    ("Pruebas de Carga", "Backend", "Batch processing", "Transcripción y análisis masivo con 50+ audios."),
    ("Pruebas de Seguridad", "Full Stack", "Manual", "Acceso sin token, tokens expirados, acceso a datos de otro usuario."),
]
create_table(doc, ["Tipo", "Scope", "Herramienta", "Cobertura"], test_types, col_widths=[3.5, 2, 3.5, 7.5])

add_heading("16.2. Casos de Prueba Principales", level=2)

test_cases = [
    ("TC-001", "Auth", "Registro con username nuevo", "201 Created + UserResponse", "Alta"),
    ("TC-002", "Auth", "Registro con username duplicado", "409 Conflict", "Alta"),
    ("TC-003", "Auth", "Login con credenciales válidas", "200 + access_token", "Alta"),
    ("TC-004", "Auth", "Login con credenciales inválidas", "401 Unauthorized", "Alta"),
    ("TC-005", "Auth", "Acceso a ruta protegida sin token", "401 Unauthorized", "Alta"),
    ("TC-006", "Auth", "Acceso con token expirado", "401 + redirect a login", "Alta"),
    ("TC-007", "Campaigns", "Crear campaña con nombre válido", "201 Created", "Alta"),
    ("TC-008", "Campaigns", "Listar campañas con paginación", "200 + PageResponse", "Alta"),
    ("TC-009", "Campaigns", "Eliminar campaña con audios", "204 + CASCADE delete", "Alta"),
    ("TC-010", "Audios", "Subir archivo MP3 válido", "201 + AudioResponse", "Alta"),
    ("TC-011", "Audios", "Subir archivo con formato inválido (.exe)", "400 Bad Request", "Alta"),
    ("TC-012", "Audios", "Subir archivo mayor a 500MB", "400 Bad Request", "Alta"),
    ("TC-013", "Audios", "Carga masiva de 10 archivos", "201 + array de AudioResponse", "Alta"),
    ("TC-014", "Transcripción", "Transcribir con Deepgram", "200 + transcripción timestamped", "Alta"),
    ("TC-015", "Transcripción", "Transcripción masiva de campaña", "200 + progreso polling", "Alta"),
    ("TC-016", "Transcripción", "Detener transcripción en curso", "200 + cancelled=True", "Media"),
    ("TC-017", "Análisis", "Analizar con prompt personalizado", "200 + JSON estructurado", "Alta"),
    ("TC-018", "Análisis", "Análisis masivo asíncrono", "200 + progreso polling", "Alta"),
    ("TC-019", "Análisis", "Verificar evaluación (Cumple/No cumple/No aplica)", "Valores correctos", "Alta"),
    ("TC-020", "Pipeline", "Pipeline completo: transcribir + analizar", "200 + resultados", "Alta"),
    ("TC-021", "Export", "Exportar resultados a XLSX", "200 + archivo descargable", "Alta"),
    ("TC-022", "Export", "Exportar con transcripciones incluidas", "200 + columnas extra", "Media"),
    ("TC-023", "Dashboard", "Cargar estadísticas generales", "200 + datos correctos", "Media"),
    ("TC-024", "UI", "Cambiar tema claro/oscuro", "Persistencia en localStorage", "Baja"),
    ("TC-025", "Seguridad", "Acceder a campaña de otro usuario", "404 Not Found", "Alta"),
]
create_table(doc, ["ID", "Módulo", "Caso", "Resultado Esperado", "Prioridad"], test_cases, col_widths=[1.5, 2, 4.5, 5.5, 1.5])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 17. GLOSARIO
# ══════════════════════════════════════════════════════════════════════════════

add_heading("17. GLOSARIO DE TÉRMINOS", level=1)

glossary = [
    ("API", "Application Programming Interface. Interfaz de comunicación entre sistemas."),
    ("ASGI", "Asynchronous Server Gateway Interface. Interfaz de servidor para aplicaciones Python asíncronas."),
    ("Batch Processing", "Procesamiento por lotes. Ejecución de múltiples operaciones como grupo."),
    ("Bearer Token", "Token de autorización enviado en cabecera HTTP para identificar al usuario."),
    ("CORS", "Cross-Origin Resource Sharing. Mecanismo de seguridad para peticiones entre dominios."),
    ("CRUD", "Create, Read, Update, Delete. Operaciones básicas de persistencia."),
    ("Diarización", "Proceso de identificar y separar diferentes hablantes en una grabación de audio."),
    ("DTO", "Data Transfer Object. Objeto para transferir datos entre capas del sistema."),
    ("FastAPI", "Framework web Python moderno, de alto rendimiento, para construir APIs."),
    ("JWT", "JSON Web Token. Estándar para tokens de autenticación seguros."),
    ("Lazy Loading", "Carga diferida. Técnica para cargar módulos solo cuando se necesitan."),
    ("ORM", "Object-Relational Mapping. Técnica para convertir entre objetos y registros de BD."),
    ("Pipeline", "Secuencia de procesos encadenados donde la salida de uno es la entrada del siguiente."),
    ("Prompt", "Instrucción de texto que se envía al modelo de IA para obtener una respuesta específica."),
    ("QA", "Quality Assurance. Garantía de calidad. Evaluación de cumplimiento de estándares."),
    ("REST", "Representational State Transfer. Estilo de arquitectura para APIs web."),
    ("Signal", "Primitiva reactiva de Angular para gestión de estado sincrónico."),
    ("SPA", "Single Page Application. Aplicación web que carga una sola página y actualiza dinámicamente."),
    ("Standalone Component", "Componente Angular autónomo que no requiere NgModule."),
    ("Transcripción", "Conversión de audio a texto. Incluye timestamps y etiquetas de hablante."),
    ("UUID", "Universally Unique Identifier. Identificador único de 128 bits."),
    ("WebSocket", "Protocolo de comunicación bidireccional en tiempo real (futuro)."),
    ("WhisperX", "Modelo de reconocimiento de voz con alineación temporal y soporte multi-idioma."),
]
create_table(doc, ["Término", "Definición"], glossary, col_widths=[4, 12.5])

# ══════════════════════════════════════════════════════════════════════════════
# GUARDAR DOCUMENTO
# ══════════════════════════════════════════════════════════════════════════════

output_path = os.path.join(
    r"C:\Users\Luis\Documents\.repositorios_github",
    "METODOLOGIA_DESARROLLO_TRANSCRIPTION_ANALYSIS.docx"
)
doc.save(output_path)
print(f"Documento generado exitosamente: {output_path}")
